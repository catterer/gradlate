#!/usr/bin/env python3
#coding: utf8
import sys
import os
import pickle
import getopt
import re
from nltk.translate.gale_church import align_blocks,align_texts
from nltk.translate.api import AlignedSent
import nltk.data
from docx import Document

block_separator = re.compile('\n\n\n\n[\n\r]*')
stnc_tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')
chapter_re = re.compile('[\r\n]*(CHAPTER *[A-Z]*) *$', re.DOTALL)
part_re = re.compile('[\r\n]*(PART *[A-Z]*) *$')
min_parag_len = 10

def help_exit():
    print('test.py -f <from_translation> -t <to_translation> -o <outputfile>')
    sys.exit(2)

class SentSep:
    def split(s):
        tokens = [tok for tok in sentence_separator.split() if tok != '']


class Sentence:
    def __init__(self, stc_raw, is_header=0):
        self.is_header = is_header
        self.raw = stc_raw.replace('\n', ' ').replace('\r', '')
    def __repr__(self):
        if self.is_header:
            return '<{}>{}</>'.format(self.is_header, self.raw)
        return self.raw

class Block:
    def __init__(self, block_raw, debug=None):
        self.raw = block_raw
        self.sentences = []
        for s in stnc_tokenizer.tokenize(self.raw):
            new_sents = []
            m = part_re.match(s)
            if m:
                new_sents = [Sentence(m.group(1), is_header=1)]
            else:
                m = chapter_re.match(s)
                if m:
                    new_sents = [Sentence(m.group(1), is_header=2)]
                else:
                    new_sents = [Sentence(s)]

            self.sentences += new_sents
            if debug:
                debug('"{}"\n=>{}\n'.format(s, new_sents))

        self.stnc_lengths_char = [len(s.raw) for s in self.sentences]

class Text:
    def __init__(self, filename, separ = block_separator, debug_output=None):
        with open(filename, 'r') as f:
            self.raw = f.read()
        debug = None
        if debug_output:
            with open(debug_output, 'w') as fd:
                self.blocks = [Block(b, debug=fd.write) for b in separ.split(self.raw)]
        else:
            self.blocks = [Block(b) for b in separ.split(self.raw)]

def chunk_list(seq, num):
    avg = len(seq) / float(num)
    out = []
    last = 0.0

    while last < len(seq):
        out.append(seq[int(last):int(last + avg)])
        last += avg

    return out

class TextXn:
    def __init__(self, text_f, text_t, logger=print):
        if len(text_f.blocks) != len(text_t.blocks):
            raise Exception('different amount of blocks in texts')

        self.logger = logger
        self.text_f = text_f
        self.text_t = text_t
        self.aligned_blocks = {}
        self.bitex = []

    def blocks_number(self):
        return len(self.text_f.blocks)

    def align_block(self, block_n):
        fb = self.text_f.blocks[block_n]
        tb = self.text_t.blocks[block_n]
        self.aligned_blocks[block_n] = align_blocks(fb.stnc_lengths_char, tb.stnc_lengths_char)

    def build_bitex(self):
        for i in range(0, self.blocks_number()):
            self.logger('aligning block {}'.format(i))
            self.align_block(i)

        self.bitex = []
        for b_id in self.aligned_blocks.keys():
            fid_last = None
            tid_last = None

            for (fid, tid) in self.aligned_blocks[b_id]:
                assert ((fid, tid) != (fid_last, tid_last))
                fsn = self.text_f.blocks[b_id].sentences[fid]
                tsn = self.text_t.blocks[b_id].sentences[tid]
                if fid != fid_last and tid != tid_last:
                    self.bitex.append((fsn, tsn))
                elif fid == fid_last:
                    self.bitex[-1][1].raw += tsn.raw
                else:
                    self.bitex[-1][0].raw += fsn.raw
                (fid_last,tid_last) = (fid, tid)

    def dump(self, fname):
        with open(fname, 'wb') as fd:
            pickle.dump(self, fd)

    def form_bilingual_table(self, fname):
        d = Document()
        table = d.add_table(rows = 0, cols = 2)
        for (f, t) in self.bitex:
            cells = table.add_row().cells
            cells[0].text = f.raw
            cells[1].text = t.raw
        d.save(fname)

    def form_bilingual_doc(self, fname):
        d = Document()
        f_accum = ''
        t_accum = ''

        def flush():
            nonlocal f_accum
            nonlocal t_accum
            d.add_paragraph(f_accum, style='Quote')
            d.add_paragraph(t_accum)
            f_accum = ''
            t_accum = ''

        for (f, t) in self.bitex:
            if f.is_header:
                flush()
                d.add_heading(f.raw, f.is_header)
                continue;
            f_accum = '{} {}'.format(f_accum, f.raw)
            t_accum = '{} {}'.format(t_accum, t.raw)
            if len(f_accum) > min_parag_len and len(t_accum) > min_parag_len:
                flush()
        flush()
        d.save(fname)

    def form_bilingual_text(self, fname):
        with open(fname, 'w') as fd:
            for (f, t) in self.bitex:
                if f.is_header:
                    fd.write('{} {}\n'.format(f.is_header, f.raw))
                    continue;
                fd.write('{}\n'.format(f.raw))
                fd.write('  {}\n'.format(t.raw))


def build_debug(d):
    l = os.listdir(d)
    for f in l:
        Text(os.path.join(d, f), debug_output='./dbg_{}'.format(f))


if __name__ == '__main__':
    try:
        opts, _ = getopt.getopt(sys.argv[1:],'hf:t:o:x:',['from=','to=','out=','xn='])
    except getopt.GetoptError:
        help_exit()

    from_path = None
    to_path = None
    out_path = None
    xn_path = None

    for k,v in opts:
        if k == '-h': help_exit()
        if k == '-f': from_path = v
        if k == '-t': to_path = v
        if k == '-o': out_path = v
        if k == '-x': xn_path = v

    if not xn_path and not all([from_path, to_path]):
        help_exit()

    if xn_path:
        with open(xn_path, 'rb') as xn_file:
            xn = pickle.load(xn_file)
    else:
        from_text = Text(from_path)
        to_text = Text(to_path)
        xn = TextXn(from_text, to_text)
        xn.build_bitex()
        xn.dump('.model')

    xn.form_bilingual_doc(out_path)



